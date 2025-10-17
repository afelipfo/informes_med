"""
Módulo de generación de reportes para Survey123
Secretaría de Infraestructura Física de Medellín

Este módulo proporciona funcionalidades para generar reportes
en diferentes formatos (PDF, Word, Excel) a partir de los datos procesados.
"""

import pandas as pd
import numpy as np
from datetime import datetime, timedelta
import os
import json
from typing import Dict, List, Any, Optional
import base64
import io

# Importaciones para generación de reportes
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    REPORTLAB_DISPONIBLE = True
except ImportError:
    REPORTLAB_DISPONIBLE = False
    print("ADVERTENCIA: ReportLab no está instalado. Instale con: pip install reportlab")

try:
    from docx import Document
    from docx.shared import Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_DISPONIBLE = True
except ImportError:
    DOCX_DISPONIBLE = False
    print("ADVERTENCIA: python-docx no está instalado. Instale con: pip install python-docx")

import matplotlib.pyplot as plt
import seaborn as sns

# Importar plotly si está disponible
try:
    import plotly.express as px
    import plotly.graph_objects as go
    PLOTLY_DISPONIBLE = True
except ImportError:
    PLOTLY_DISPONIBLE = False
    px = None
    go = None

class GeneradorReportes:
    """
    Clase principal para generar reportes de Survey123
    """
    
    def __init__(self, datos: pd.DataFrame, metricas: Dict[str, Any] = None):
        """
        Inicializa el generador de reportes
        
        Args:
            datos: DataFrame con los datos procesados
            metricas: Dict con métricas calculadas (opcional)
        """
        self.datos = datos.copy()
        self.metricas = metricas or {}
        self.fecha_reporte = datetime.now()
        self.configurar_estilos()
    
    def configurar_estilos(self):
        """Configura estilos para los reportes"""
        self.colores_institucionales = {
            'azul_principal': '#003366',
            'azul_secundario': '#0066CC',
            'verde_acento': '#00CC66',
            'gris_texto': '#333333',
            'gris_claro': '#F5F5F5'
        }
        
        # Configurar matplotlib para español
        plt.rcParams['font.size'] = 10
        plt.rcParams['axes.titlesize'] = 12
        plt.rcParams['figure.figsize'] = (10, 6)
    
    def generar_portada_datos(self) -> Dict[str, str]:
        """
        Genera los datos para la portada del reporte
        
        Returns:
            Dict con información de la portada
        """
        return {
            'titulo': 'INFORME DE SEGUIMIENTO',
            'subtitulo': 'OBRAS DE INFRAESTRUCTURA FÍSICA',
            'entidad': 'SECRETARÍA DE INFRAESTRUCTURA FÍSICA',
            'ciudad': 'ALCALDÍA DE MEDELLÍN',
            'fecha': self.fecha_reporte.strftime('%d de %B de %Y'),
            'periodo_analisis': f"{self.datos['fecha_dilig'].min().strftime('%d/%m/%Y')} - {self.datos['fecha_dilig'].max().strftime('%d/%m/%Y')}" if 'fecha_dilig' in self.datos.columns else 'N/A',
            'total_registros': len(self.datos),
            'version': '1.0'
        }
    
    def generar_resumen_ejecutivo(self) -> Dict[str, Any]:
        """
        Genera el resumen ejecutivo del reporte
        
        Returns:
            Dict con el resumen ejecutivo
        """
        resumen = {
            'total_obras': len(self.datos),
            'estados_obra': self.datos['estado_obr'].value_counts().to_dict() if 'estado_obr' in self.datos.columns else {},
            'duracion_proyecto': (self.datos['fecha_dilig'].max() - self.datos['fecha_dilig'].min()).days if 'fecha_dilig' in self.datos.columns else 0,
            'cobertura_geografica': {
                'puntos_unicos': self.datos['id_punto'].nunique() if 'id_punto' in self.datos.columns else 0,
                'rango_latitud': f"{self.datos['Y'].min():.4f} - {self.datos['Y'].max():.4f}" if 'Y' in self.datos.columns else 'N/A',
                'rango_longitud': f"{self.datos['X'].min():.4f} - {self.datos['X'].max():.4f}" if 'X' in self.datos.columns else 'N/A'
            }
        }
        
        # Agregar métricas de recursos humanos si están disponibles
        if 'cant_ayuda' in self.datos.columns and 'cant_ofici' in self.datos.columns:
            resumen['recursos_humanos'] = {
                'total_personal': self.datos['cant_ayuda'].sum() + self.datos['cant_ofici'].sum(),
                'promedio_por_obra': (self.datos['cant_ayuda'] + self.datos['cant_ofici']).mean()
            }
        
        return resumen
    
    def crear_graficos_reporte(self, ruta_salida: str = 'datos/reportes_generados/graficos/') -> Dict[str, str]:
        """
        Crea los gráficos para incluir en el reporte
        
        Args:
            ruta_salida: Ruta donde guardar los gráficos
            
        Returns:
            Dict con las rutas de los gráficos generados
        """
        os.makedirs(ruta_salida, exist_ok=True)
        graficos = {}
        
        # Gráfico 1: Estados de obra
        if 'estado_obr' in self.datos.columns:
            plt.figure(figsize=(8, 6))
            estados = self.datos['estado_obr'].value_counts()
            plt.pie(estados.values, labels=estados.index, autopct='%1.1f%%',
                   colors=['#003366', '#0066CC', '#00CC66', '#FF6600'])
            plt.title('Distribución por Estado de Obra', fontsize=14, fontweight='bold')
            grafico_estados = f'{ruta_salida}estados_obra.png'
            plt.savefig(grafico_estados, dpi=300, bbox_inches='tight')
            plt.close()
            graficos['estados_obra'] = grafico_estados
        
        # Gráfico 2: Recursos humanos por obra
        if 'cant_ayuda' in self.datos.columns and 'cant_ofici' in self.datos.columns:
            plt.figure(figsize=(10, 6))
            x = range(min(20, len(self.datos)))  # Máximo 20 obras para legibilidad
            ayudantes = self.datos['cant_ayuda'].iloc[:20]
            oficiales = self.datos['cant_ofici'].iloc[:20]
            
            width = 0.35
            plt.bar([i - width/2 for i in x], ayudantes, width, label='Ayudantes', color='#0066CC')
            plt.bar([i + width/2 for i in x], oficiales, width, label='Oficiales', color='#003366')
            
            plt.xlabel('Número de Obra')
            plt.ylabel('Cantidad de Personal')
            plt.title('Recursos Humanos por Obra (Primeras 20 obras)', fontweight='bold')
            plt.legend()
            plt.grid(True, alpha=0.3)
            
            grafico_rrhh = f'{ruta_salida}recursos_humanos.png'
            plt.savefig(grafico_rrhh, dpi=300, bbox_inches='tight')
            plt.close()
            graficos['recursos_humanos'] = grafico_rrhh
        
        # Gráfico 3: Línea de tiempo de obras
        if 'fecha_dilig' in self.datos.columns:
            plt.figure(figsize=(12, 6))
            obras_por_fecha = self.datos.groupby('fecha_dilig').size()
            plt.plot(obras_por_fecha.index, obras_por_fecha.values, 
                    marker='o', linewidth=2, markersize=4, color='#003366')
            plt.fill_between(obras_por_fecha.index, obras_por_fecha.values, alpha=0.3, color='#0066CC')
            
            plt.xlabel('Fecha')
            plt.ylabel('Número de Obras')
            plt.title('Cronograma de Obras por Fecha', fontweight='bold')
            plt.xticks(rotation=45)
            plt.grid(True, alpha=0.3)
            
            grafico_cronograma = f'{ruta_salida}cronograma_obras.png'
            plt.savefig(grafico_cronograma, dpi=300, bbox_inches='tight')
            plt.close()
            graficos['cronograma'] = grafico_cronograma
        
        return graficos
    
    def generar_reporte_pdf(self, nombre_archivo: str = None, incluir_graficos: bool = True) -> str:
        """
        Genera un reporte en formato PDF

        Args:
            nombre_archivo: Nombre del archivo de salida
            incluir_graficos: Si incluir gráficos en el reporte

        Returns:
            Ruta del archivo generado

        Raises:
            ImportError: Si ReportLab no está instalado
        """
        if not REPORTLAB_DISPONIBLE:
            raise ImportError("ReportLab no está instalado. Instale con: pip install reportlab")

        return self._generar_reporte_pdf_interno(nombre_archivo, incluir_graficos)

    def _generar_reporte_pdf_interno(self, nombre_archivo: str = None, incluir_graficos: bool = True) -> str:
        """
        Genera un reporte en formato PDF
        
        Args:
            nombre_archivo: Nombre del archivo de salida
            incluir_graficos: Si incluir gráficos en el reporte
            
        Returns:
            Ruta del archivo generado
        """
        if nombre_archivo is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            nombre_archivo = f'informe_survey123_{timestamp}.pdf'
        
        ruta_salida = f'datos/reportes_generados/{nombre_archivo}'
        os.makedirs('datos/reportes_generados/', exist_ok=True)
        
        # Crear documento PDF
        doc = SimpleDocTemplate(ruta_salida, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # Estilo personalizado para títulos
        titulo_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#003366'),
            alignment=TA_CENTER,
            spaceAfter=30
        )
        
        # Portada
        portada = self.generar_portada_datos()
        story.append(Paragraph(portada['titulo'], titulo_style))
        story.append(Paragraph(portada['subtitulo'], titulo_style))
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph(portada['entidad'], styles['Heading2']))
        story.append(Paragraph(portada['ciudad'], styles['Heading3']))
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph(f"Fecha: {portada['fecha']}", styles['Normal']))
        story.append(Paragraph(f"Período de análisis: {portada['periodo_analisis']}", styles['Normal']))
        story.append(Spacer(1, 1*inch))
        
        # Resumen ejecutivo
        story.append(Paragraph("RESUMEN EJECUTIVO", styles['Heading1']))
        resumen = self.generar_resumen_ejecutivo()
        
        story.append(Paragraph(f"Total de obras registradas: {resumen['total_obras']}", styles['Normal']))
        story.append(Paragraph(f"Duración del proyecto: {resumen['duracion_proyecto']} días", styles['Normal']))
        
        # Tabla de estados de obra
        if resumen['estados_obra']:
            story.append(Paragraph("Estados de Obra:", styles['Heading3']))
            tabla_estados = [['Estado', 'Cantidad', 'Porcentaje']]
            total = sum(resumen['estados_obra'].values())
            for estado, cantidad in resumen['estados_obra'].items():
                porcentaje = (cantidad / total) * 100
                tabla_estados.append([estado, str(cantidad), f"{porcentaje:.1f}%"])
            
            t = Table(tabla_estados)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#003366')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F5F5F5')),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(t)
            story.append(Spacer(1, 0.3*inch))
        
        # Incluir gráficos si se solicita
        if incluir_graficos:
            graficos = self.crear_graficos_reporte()
            for nombre, ruta in graficos.items():
                if os.path.exists(ruta):
                    story.append(Paragraph(f"Gráfico: {nombre.replace('_', ' ').title()}", styles['Heading3']))
                    img = Image(ruta, width=6*inch, height=4*inch)
                    story.append(img)
                    story.append(Spacer(1, 0.2*inch))
        
        # Construir PDF
        doc.build(story)
        return ruta_salida
    
    def generar_reporte_word(self, nombre_archivo: str = None) -> str:
        """
        Genera un reporte en formato Word

        Args:
            nombre_archivo: Nombre del archivo de salida

        Returns:
            Ruta del archivo generado

        Raises:
            ImportError: Si python-docx no está instalado
        """
        if not DOCX_DISPONIBLE:
            raise ImportError("python-docx no está instalado. Instale con: pip install python-docx")

        return self._generar_reporte_word_interno(nombre_archivo)

    def _generar_reporte_word_interno(self, nombre_archivo: str = None) -> str:
        """
        Genera un reporte en formato Word
        
        Args:
            nombre_archivo: Nombre del archivo de salida
            
        Returns:
            Ruta del archivo generado
        """
        if nombre_archivo is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            nombre_archivo = f'informe_survey123_{timestamp}.docx'
        
        ruta_salida = f'datos/reportes_generados/{nombre_archivo}'
        os.makedirs('datos/reportes_generados/', exist_ok=True)
        
        # Crear documento Word
        doc = Document()
        
        # Portada
        portada = self.generar_portada_datos()
        titulo = doc.add_heading(portada['titulo'], 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitulo = doc.add_heading(portada['subtitulo'], 1)
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        entidad = doc.add_heading(portada['entidad'], 2)
        entidad.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        ciudad = doc.add_heading(portada['ciudad'], 3)
        ciudad.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph(f"Fecha: {portada['fecha']}")
        doc.add_paragraph(f"Período de análisis: {portada['periodo_analisis']}")
        doc.add_paragraph(f"Total de registros: {portada['total_registros']}")
        
        # Salto de página
        doc.add_page_break()
        
        # Resumen ejecutivo
        doc.add_heading('RESUMEN EJECUTIVO', 1)
        resumen = self.generar_resumen_ejecutivo()
        
        doc.add_paragraph(f"El presente informe analiza {resumen['total_obras']} obras de infraestructura "
                         f"física registradas durante un período de {resumen['duracion_proyecto']} días.")
        
        # Estados de obra
        if resumen['estados_obra']:
            doc.add_heading('Estados de Obra', 2)
            tabla = doc.add_table(rows=1, cols=3)
            tabla.style = 'Table Grid'
            tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Encabezados
            hdr_cells = tabla.rows[0].cells
            hdr_cells[0].text = 'Estado'
            hdr_cells[1].text = 'Cantidad'
            hdr_cells[2].text = 'Porcentaje'
            
            # Datos
            total = sum(resumen['estados_obra'].values())
            for estado, cantidad in resumen['estados_obra'].items():
                row_cells = tabla.add_row().cells
                row_cells[0].text = estado
                row_cells[1].text = str(cantidad)
                row_cells[2].text = f"{(cantidad/total)*100:.1f}%"
        
        # Recursos humanos
        if 'recursos_humanos' in resumen:
            doc.add_heading('Recursos Humanos', 2)
            doc.add_paragraph(f"Total de personal asignado: {resumen['recursos_humanos']['total_personal']}")
            doc.add_paragraph(f"Promedio de personal por obra: {resumen['recursos_humanos']['promedio_por_obra']:.1f}")
        
        # Cobertura geográfica
        doc.add_heading('Cobertura Geográfica', 2)
        doc.add_paragraph(f"Puntos únicos de intervención: {resumen['cobertura_geografica']['puntos_unicos']}")
        doc.add_paragraph(f"Rango de latitud: {resumen['cobertura_geografica']['rango_latitud']}")
        doc.add_paragraph(f"Rango de longitud: {resumen['cobertura_geografica']['rango_longitud']}")
        
        # Guardar documento
        doc.save(ruta_salida)
        return ruta_salida
    
    def generar_reporte_excel(self, nombre_archivo: str = None) -> str:
        """
        Genera un reporte en formato Excel con múltiples hojas
        
        Args:
            nombre_archivo: Nombre del archivo de salida
            
        Returns:
            Ruta del archivo generado
        """
        if nombre_archivo is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            nombre_archivo = f'informe_survey123_{timestamp}.xlsx'
        
        ruta_salida = f'datos/reportes_generados/{nombre_archivo}'
        os.makedirs('datos/reportes_generados/', exist_ok=True)
        
        with pd.ExcelWriter(ruta_salida, engine='openpyxl') as writer:
            # Hoja 1: Datos originales
            self.datos.to_excel(writer, sheet_name='Datos_Originales', index=False)
            
            # Hoja 2: Resumen ejecutivo
            resumen = self.generar_resumen_ejecutivo()
            df_resumen = pd.DataFrame([
                ['Total de obras', resumen['total_obras']],
                ['Duración del proyecto (días)', resumen['duracion_proyecto']],
                ['Puntos únicos', resumen['cobertura_geografica']['puntos_unicos']]
            ], columns=['Métrica', 'Valor'])
            df_resumen.to_excel(writer, sheet_name='Resumen_Ejecutivo', index=False)
            
            # Hoja 3: Estados de obra
            if resumen['estados_obra']:
                df_estados = pd.DataFrame(list(resumen['estados_obra'].items()), 
                                        columns=['Estado', 'Cantidad'])
                df_estados['Porcentaje'] = (df_estados['Cantidad'] / df_estados['Cantidad'].sum()) * 100
                df_estados.to_excel(writer, sheet_name='Estados_Obra', index=False)
            
            # Hoja 4: Análisis temporal
            if 'fecha_dilig' in self.datos.columns:
                df_temporal = self.datos.groupby('fecha_dilig').size().reset_index()
                df_temporal.columns = ['Fecha', 'Numero_Obras']
                df_temporal.to_excel(writer, sheet_name='Analisis_Temporal', index=False)
            
            # Hoja 5: Recursos humanos (si disponible)
            if 'cant_ayuda' in self.datos.columns and 'cant_ofici' in self.datos.columns:
                df_rrhh = self.datos[['id_punto', 'cant_ayuda', 'cant_ofici']].copy()
                df_rrhh['total_personal'] = df_rrhh['cant_ayuda'] + df_rrhh['cant_ofici']
                df_rrhh.to_excel(writer, sheet_name='Recursos_Humanos', index=False)
        
        return ruta_salida
    
    def generar_reporte_completo(self, formato: str = 'todos', prefijo_nombre: str = None) -> Dict[str, str]:
        """
        Genera reportes completos en todos los formatos solicitados
        
        Args:
            formato: 'pdf', 'word', 'excel' o 'todos'
            prefijo_nombre: Prefijo para los nombres de archivo
            
        Returns:
            Dict con las rutas de los archivos generados
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        prefijo = prefijo_nombre or 'informe_survey123'
        
        archivos_generados = {}
        
        if formato in ['pdf', 'todos']:
            try:
                archivo_pdf = self.generar_reporte_pdf(f'{prefijo}_{timestamp}.pdf')
                archivos_generados['pdf'] = archivo_pdf
            except Exception as e:
                print(f"Error generando PDF: {e}")
        
        if formato in ['word', 'todos']:
            try:
                archivo_word = self.generar_reporte_word(f'{prefijo}_{timestamp}.docx')
                archivos_generados['word'] = archivo_word
            except Exception as e:
                print(f"Error generando Word: {e}")
        
        if formato in ['excel', 'todos']:
            try:
                archivo_excel = self.generar_reporte_excel(f'{prefijo}_{timestamp}.xlsx')
                archivos_generados['excel'] = archivo_excel
            except Exception as e:
                print(f"Error generando Excel: {e}")
        
        return archivos_generados

def generar_reporte_automatico(datos: pd.DataFrame, metricas: Dict[str, Any] = None,
                             formato: str = 'todos') -> Dict[str, str]:
    """
    Función principal para generar reportes automáticamente
    
    Args:
        datos: DataFrame con los datos procesados
        metricas: Dict con métricas calculadas
        formato: Formato de salida ('pdf', 'word', 'excel', 'todos')
        
    Returns:
        Dict con las rutas de los archivos generados
    """
    generador = GeneradorReportes(datos, metricas)
    return generador.generar_reporte_completo(formato)

if __name__ == "__main__":
    print("Módulo de generación de reportes Survey123 - Secretaría de Infraestructura Física de Medellín")
    print("Para usar este módulo, importe la clase GeneradorReportes y pase sus datos procesados.")
